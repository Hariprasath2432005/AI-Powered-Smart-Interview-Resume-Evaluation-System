# ============================================================
# CAREERAI - AI POWERED RESUME & INTERVIEW ASSISTANT
# ============================================================

import streamlit as st
import os
import re

from dotenv import load_dotenv

from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
from langchain_groq import ChatGroq

from pypdf import PdfReader
import docx


# ============================================================
# 1. LOAD ENVIRONMENT
# ============================================================

load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")


# ============================================================
# 2. INITIALIZE LLM
# ============================================================

llm = ChatGroq(
    groq_api_key=GROQ_API_KEY,
    model_name="llama-3.1-8b-instant"
)


# ============================================================
# 3. EMBEDDING MODEL
# ============================================================

embedding_model = SentenceTransformer(
    "all-MiniLM-L6-v2"
)


# ============================================================
# 4. PAGE CONFIG
# ============================================================

st.set_page_config(
    page_title="CareerAI",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)


# ============================================================
# 5. SESSION STATE
# ============================================================

if "basic_questions" not in st.session_state:
    st.session_state.basic_questions = None

if "personalized_questions" not in st.session_state:
    st.session_state.personalized_questions = None

if "resume_analysis" not in st.session_state:
    st.session_state.resume_analysis = None

if "job_analysis" not in st.session_state:
    st.session_state.job_analysis = None

if "resume_improvement" not in st.session_state:
    st.session_state.resume_improvement = None

if "project_analysis" not in st.session_state:
    st.session_state.project_analysis = None

if "job_risk_analysis" not in st.session_state:
    st.session_state.job_risk_analysis = None


# ============================================================
# 6. CUSTOM CSS
# ============================================================

st.markdown(
    """
<style>

.stApp {
    background: #0f172a;
}

.main {
    background: #0f172a;
}

.block-container {
    max-width: 1250px;
    padding-top: 28px;
    padding-bottom: 60px;
}


/* SIDEBAR */

section[data-testid="stSidebar"] {
    background: #111827;
    border-right: 1px solid #263244;
}

section[data-testid="stSidebar"] > div {
    padding-top: 25px;
}

section[data-testid="stSidebar"] * {
    color: #e5e7eb;
}

section[data-testid="stSidebar"] label {
    color: #e5e7eb !important;
    font-weight: 600 !important;
}

section[data-testid="stSidebar"] input,
section[data-testid="stSidebar"] textarea {
    background: #1e293b !important;
    color: #f8fafc !important;
    border: 1px solid #334155 !important;
    border-radius: 10px !important;
}

section[data-testid="stSidebar"] input::placeholder,
section[data-testid="stSidebar"] textarea::placeholder {
    color: #94a3b8 !important;
}


/* BRAND */

.brand-title {
    font-size: 30px;
    font-weight: 800;
    color: #f8fafc;
    letter-spacing: -0.5px;
    margin-bottom: 3px;
}

.brand-title span {
    color: #818cf8;
}

.brand-subtitle {
    color: #94a3b8;
    font-size: 13px;
    margin-bottom: 20px;
}


/* HERO */

.hero-box {
    background: linear-gradient(
        135deg,
        #172033 0%,
        #1e293b 55%,
        #25204a 100%
    );

    border: 1px solid #334155;
    border-radius: 20px;

    padding: 32px 38px;

    margin-bottom: 30px;

    box-shadow:
        0 15px 35px rgba(0,0,0,0.20);
}

.hero-small {
    color: #a5b4fc;
    font-size: 12px;
    font-weight: 700;
    letter-spacing: 1.5px;
    text-transform: uppercase;
    margin-bottom: 10px;
}

.hero-title {
    color: #f8fafc;
    font-size: 38px;
    font-weight: 800;
    line-height: 1.2;
    margin: 0;
    letter-spacing: -1px;
}

.hero-title span {
    color: #818cf8;
}

.hero-description {
    color: #cbd5e1;
    font-size: 15px;
    line-height: 1.7;
    margin-top: 12px;
    max-width: 850px;
}


/* FEATURE BADGES */

.feature-row {
    display: flex;
    flex-wrap: wrap;
    gap: 9px;
    margin-top: 20px;
}

.feature-badge {
    background: #26324a;
    border: 1px solid #3b4963;
    color: #e2e8f0;
    padding: 7px 12px;
    border-radius: 20px;
    font-size: 12px;
    font-weight: 600;
}


/* SECTION */

.section-title {
    color: #f8fafc;
    font-size: 28px;
    font-weight: 800;
    margin-top: 5px;
    margin-bottom: 5px;
}

.section-title span {
    color: #818cf8;
}

.section-description {
    color: #94a3b8;
    font-size: 14px;
    margin-bottom: 22px;
}


/* DASHBOARD */

.dashboard-heading {
    margin-top: 5px;
    margin-bottom: 18px;
}

.dashboard-card {
    background: #172033;
    border: 1px solid #2d3a50;
    border-radius: 15px;

    padding: 22px;

    height: 145px;
    min-height: 145px;

    display: flex;
    flex-direction: column;
    justify-content: space-between;

    box-sizing: border-box;
}

.dashboard-icon {
    font-size: 20px;
    margin-bottom: 7px;
}

.dashboard-title {
    color: #94a3b8;
    font-size: 11px;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 0.7px;
}

.dashboard-value {
    color: #f8fafc;
    font-size: 14px;
    font-weight: 700;
    word-break: break-word;
}

.dashboard-bottom-space {
    height: 18px;
}


/* TABS */

.stTabs {
    margin-top: 8px;
}

.stTabs [data-baseweb="tab-list"] {
    gap: 5px;

    background: #172033;

    padding: 6px;

    border: 1px solid #2d3a50;

    border-radius: 12px;

    margin-bottom: 30px;
}

.stTabs [data-baseweb="tab"] {
    color: #94a3b8 !important;

    background: transparent !important;

    border-radius: 8px;

    padding: 10px 15px;

    font-size: 13px;

    font-weight: 650;
}

.stTabs [data-baseweb="tab"]:hover {
    color: #c7d2fe !important;
    background: #222d42 !important;
}

.stTabs [aria-selected="true"] {
    color: #ffffff !important;
    background: #4f46e5 !important;
}

.stTabs [aria-selected="true"] * {
    color: #ffffff !important;
}


/* CARD */

.card {
    background: #172033;

    border: 1px solid #2d3a50;

    border-radius: 16px;

    padding: 22px;

    margin-bottom: 18px;

    box-sizing: border-box;
}

.card-title {
    color: #f8fafc;

    font-size: 17px;

    font-weight: 750;

    margin-bottom: 7px;
}

.card-text {
    color: #94a3b8;

    font-size: 14px;

    line-height: 1.6;
}


/* QUESTION CARD */

.question-card {
    background: #172033;

    border: 1px solid #2d3a50;

    border-radius: 16px;

    padding: 22px;

    min-height: 220px;

    margin-bottom: 20px;

    box-sizing: border-box;
}

.question-card-title {
    color: #f8fafc;

    font-size: 18px;

    font-weight: 750;

    margin-bottom: 8px;
}

.question-card-description {
    color: #94a3b8;

    font-size: 13px;

    line-height: 1.6;

    min-height: 45px;
}


/* GENERATED BOX */

.generated-box {
    background: #111827;

    border: 1px solid #334155;

    border-radius: 12px;

    padding: 20px;

    margin-top: 12px;

    color: #e2e8f0;

    line-height: 1.8;

    font-size: 14px;
}


/* SCORE */

.score-card {
    background: #172033;

    border: 1px solid #334155;

    border-radius: 16px;

    padding: 22px;

    text-align: center;
}

.score-number {
    color: #818cf8;

    font-size: 35px;

    font-weight: 850;

    line-height: 1.2;
}

.score-label {
    color: #94a3b8;

    font-size: 13px;

    font-weight: 600;

    margin-top: 6px;
}


/* SKILLS */

.skill-container {
    display: flex;

    flex-wrap: wrap;

    gap: 8px;

    margin-top: 8px;
}

.skill {
    background: #222d49;

    border: 1px solid #3b4a70;

    color: #c7d2fe;

    padding: 7px 12px;

    border-radius: 20px;

    font-size: 12px;

    font-weight: 650;
}

.missing-skill {
    background: #3a2029;

    border: 1px solid #6b3445;

    color: #fda4af;

    padding: 7px 12px;

    border-radius: 20px;

    font-size: 12px;

    font-weight: 650;
}


/* INFO */

.info-box {
    background: #172b46;

    border: 1px solid #31567d;

    border-left: 4px solid #60a5fa;

    border-radius: 9px;

    padding: 14px 17px;

    color: #bfdbfe;

    font-size: 14px;

    margin-bottom: 15px;
}


/* BUTTON */

.stButton > button {
    width: 100%;

    background: #4f46e5;

    color: white;

    border: none;

    border-radius: 9px;

    padding: 10px 18px;

    font-size: 14px;

    font-weight: 700;

    min-height: 44px;
}

.stButton > button:hover {
    background: #6366f1;

    color: white;

    border: none;
}


/* INPUT */

.stTextInput input,
.stTextArea textarea {
    background: #172033 !important;

    color: #f8fafc !important;

    border: 1px solid #334155 !important;

    border-radius: 10px !important;
}

.stTextInput input::placeholder,
.stTextArea textarea::placeholder {
    color: #64748b !important;
}

.stTextInput input:focus,
.stTextArea textarea:focus {
    border-color: #6366f1 !important;

    box-shadow: 0 0 0 1px #6366f1 !important;
}

p {
    color: #cbd5e1;
}

h1, h2, h3, h4 {
    color: #f8fafc !important;
}


/* WELCOME */

.welcome-card {
    background: #172033;

    border: 1px solid #2d3a50;

    border-radius: 18px;

    padding: 25px;

    margin-bottom: 28px;
}

.welcome-title {
    color: #f8fafc;

    font-size: 20px;

    font-weight: 750;

    margin-bottom: 8px;
}

.welcome-text {
    color: #94a3b8;

    font-size: 14px;

    line-height: 1.7;
}


/* FOOTER */

.footer {
    text-align: center;

    color: #64748b;

    font-size: 12px;

    padding-top: 45px;

    padding-bottom: 20px;
}

hr {
    border-color: #263244 !important;
}

</style>
""",
    unsafe_allow_html=True
)


# ============================================================
# 7. SIDEBAR
# ============================================================

with st.sidebar:

    st.markdown(
        '<div class="brand-title">Career<span>AI</span></div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="brand-subtitle">'
        'Your AI-powered career assistant'
        '</div>',
        unsafe_allow_html=True
    )

    st.divider()

    st.markdown("### 📄 Resume")

    uploaded_file = st.file_uploader(
        "Upload Resume",
        type=["pdf", "docx"]
    )

    st.markdown("### 💼 Target Role")

    job_role = st.text_input(
        "Job Role",
        placeholder="e.g. AI Engineer"
    )

    st.markdown("### 📋 Job Description")

    job_description = st.text_area(
        "Paste Job Description",
        height=200,
        placeholder="Paste the complete job description here..."
    )

    st.divider()

    st.markdown("### ⚡ How it works")

    st.markdown(
        """
        **01** → Upload Resume

        **02** → Enter Target Role

        **03** → Add Job Description

        **04** → Analyze Job Match

        **05** → Prepare for Interview
        """
    )


# ============================================================
# 8. RESUME TEXT EXTRACTION
# ============================================================

def extract_resume_text(uploaded_file):

    text = ""

    try:

        if uploaded_file.name.lower().endswith(".pdf"):

            pdf = PdfReader(uploaded_file)

            for page in pdf.pages:

                page_text = page.extract_text()

                if page_text:

                    text += page_text + "\n"

        elif uploaded_file.name.lower().endswith(".docx"):

            doc = docx.Document(uploaded_file)

            for para in doc.paragraphs:

                if para.text.strip():

                    text += para.text + "\n"

            for table in doc.tables:

                for row in table.rows:

                    row_text = " ".join(
                        cell.text.strip()
                        for cell in row.cells
                    )

                    if row_text:

                        text += row_text + "\n"

    except Exception as e:

        st.error(
            f"Error reading resume: {e}"
        )

    return text.strip()


# ============================================================
# 9. CLEAN TEXT
# ============================================================

def clean_text(text):

    text = text.replace(
        "\n",
        " "
    )

    text = re.sub(
        r"\s+",
        " ",
        text
    )

    return text.strip()


# ============================================================
# 10. RESUME EVALUATION
# ============================================================

def evaluate_resume(
    resume_text,
    job_role
):

    prompt = f"""

You are an expert HR recruiter and ATS resume evaluator.

TARGET JOB ROLE:

{job_role}

RESUME:

{resume_text}

Provide:

1. ATS Score out of 100
2. Overall Resume Summary
3. Key Strengths
4. Weaknesses
5. Technical Skills Found
6. Relevant Experience
7. Education Relevance
8. Improvement Suggestions
9. Missing areas for the target role

Keep the explanation clear and practical.

"""

    response = llm.invoke(
        prompt
    )

    return response.content


# ============================================================
# 11. SKILL EXTRACTION
# ============================================================

def extract_skills(text):

    prompt = f"""

You are an expert technical recruiter.

Extract important technical skills from the text.

TEXT:

{text}

Return ONLY a comma-separated list.

Example:

Python, SQL, Machine Learning, NLP, TensorFlow

Rules:

- Only return skills.
- Do not provide explanations.
- Do not invent skills.

"""

    response = llm.invoke(
        prompt
    )

    skills_text = response.content.strip()

    skills_text = skills_text.replace(
        "\n",
        " "
    )

    skills = [
        skill.strip()
        for skill in skills_text.split(",")
        if skill.strip()
    ]

    return skills


# ============================================================
# 12. JOB COMPATIBILITY
# ============================================================

def compare_resume_with_job(
    resume_text,
    job_description
):

    resume_clean = clean_text(
        resume_text
    )

    job_clean = clean_text(
        job_description
    )

    resume_embedding = embedding_model.encode(
        [resume_clean]
    )

    job_embedding = embedding_model.encode(
        [job_clean]
    )

    similarity = cosine_similarity(
        resume_embedding,
        job_embedding
    )[0][0]

    semantic_score = round(
        max(
            0,
            min(
                1,
                similarity
            )
        ) * 100,
        2
    )

    resume_skills = extract_skills(
        resume_text
    )

    job_skills = extract_skills(
        job_description
    )

    resume_skill_lower = {
        skill.lower().strip()
        for skill in resume_skills
    }

    missing_skills = []

    for skill in job_skills:

        if skill.lower().strip() not in resume_skill_lower:

            missing_skills.append(
                skill
            )

    comparison_prompt = f"""

You are an expert technical recruiter.

Compare the candidate resume with the job description.

RESUME:

{resume_text}

JOB DESCRIPTION:

{job_description}

RESUME SKILLS:

{resume_skills}

JOB REQUIRED SKILLS:

{job_skills}

MISSING SKILLS:

{missing_skills}

Provide:

1. Matching Skills
2. Partially Matching Skills
3. Relevant Experience Match
4. Education Match
5. Strengths for This Job
6. Skill Gaps
7. Improvement Suggestions
8. Overall Hiring Recommendation

Important:

- Do not assume a skill exists if it is not in the resume.
- Keep the analysis realistic.
- Focus only on the given job description.

"""

    response = llm.invoke(
        comparison_prompt
    )

    return (
        semantic_score,
        resume_skills,
        job_skills,
        missing_skills,
        response.content
    )


# ============================================================
# 13. RESUME IMPROVEMENT
# ============================================================

def generate_resume_improvement(
    resume_text,
    job_description,
    job_role,
    missing_skills
):

    prompt = f"""

You are an expert ATS resume writer and technical recruiter.

Your job is to improve the candidate's resume specifically
for the given job description.

TARGET JOB ROLE:

{job_role}

CANDIDATE RESUME:

{resume_text}

JOB DESCRIPTION:

{job_description}

MISSING SKILLS DETECTED FROM JOB COMPATIBILITY:

{missing_skills}

IMPORTANT:

The improvement suggestions MUST be based on:

1. Job requirements
2. Missing skills
3. Missing technical capabilities
4. Missing job-specific capabilities
5. Important keywords in the job description
6. Existing projects and experience in the resume

Do NOT give generic resume advice.

Do NOT tell the candidate to add a skill if the resume does
not demonstrate that they actually know it.

Instead, explain what capability is missing and how they can
gain or demonstrate that capability.

Provide the result in this structure:

============================================================
📝 RESUME IMPROVEMENT SUGGESTIONS
============================================================

### 1. Missing Job Capabilities

List the important capabilities required by the JD that are
not clearly demonstrated in the resume.

For each:

Capability:
Why the Job Requires It:
Current Resume Gap:
How to Improve:

------------------------------------------------------------

### 2. Missing Technical Skills

For every important missing skill:

Skill:
Why It Matters:
What to Learn:
Where to Demonstrate It:

------------------------------------------------------------

### 3. Resume Section Improvements

Explain exactly what should be improved in:

- Skills
- Projects
- Experience
- Summary
- Certifications
- Education

Only mention sections where an improvement is actually needed.

------------------------------------------------------------

### 4. Project Improvement

Look at the projects actually present in the resume.

For relevant projects:

Current Project Description:
Missing Job Capability:
Recommended Improvement:

Do NOT invent project details.

Do NOT invent datasets, accuracy, tools or achievements.

------------------------------------------------------------

### 5. AI Rewritten Project Bullet Points

Rewrite existing project bullet points so that they are:

- ATS friendly
- Job relevant
- Achievement oriented
- Technically clear

IMPORTANT:

Only use information actually present in the resume.

Do not create fake achievements.

------------------------------------------------------------

### 6. Job-Specific Keywords to Add

List important keywords from the job description that can
truthfully be included in the resume.

------------------------------------------------------------

### 7. Priority Improvement Plan

Give:

🔴 High Priority
🟡 Medium Priority
🟢 Low Priority

Focus mainly on capabilities that are required by the
specific job.

FINAL RULE:

The goal is NOT to make the resume look better generally.

The goal is to make the resume demonstrate the missing
capabilities required by THIS SPECIFIC JOB DESCRIPTION.

"""

    response = llm.invoke(
        prompt
    )

    return response.content


# ============================================================
# 14. RESUME PROJECT ANALYZER
# ============================================================

def analyze_resume_projects(
    resume_text,
    job_role,
    job_description
):

    prompt = f"""

You are an expert technical recruiter and project evaluator.

Analyze ONLY the projects actually mentioned in the
candidate's resume.

CANDIDATE RESUME:

{resume_text}

TARGET JOB ROLE:

{job_role}

JOB DESCRIPTION:

{job_description}

For every project found, provide:

1. Project Name
2. Project Summary
3. Technologies / Skills Used
4. Relevance to Target Job
5. Technical Strength
6. Missing Technical Capability
7. Project Quality Score out of 100
8. How to Improve This Project
9. How This Project Can Be Explained in an Interview

Then provide:

OVERALL PROJECT PROFILE

1. Strongest Project
2. Most Relevant Project
3. Weakest Project
4. Missing Capabilities Across Projects
5. Skills demonstrated through projects
6. Skills required by JD but not demonstrated
7. Recommended project improvement priorities

IMPORTANT:

- ONLY analyze projects actually present in the resume.
- Do NOT invent projects.
- Do NOT invent technologies.
- Do NOT invent datasets.
- Do NOT invent accuracy values.
- Do NOT invent achievements.
- If something is not mentioned say:
  "Not mentioned in the resume."

"""

    response = llm.invoke(
        prompt
    )

    return response.content


# ============================================================
# 15. JOB DESCRIPTION RISK DETECTOR
# ============================================================

def analyze_job_description_risks(
    job_description
):

    prompt = f"""

You are an expert job description and recruitment analyst.

Analyze the following job description carefully.

JOB DESCRIPTION:

{job_description}

Identify important job details and potential warning signs.

============================================================
📋 JOB INFORMATION
============================================================

1. Job Title
2. Company Name
3. Location
4. Work Mode
5. Employment Type
6. Experience Required
7. Education Required
8. Salary / CTC
9. Application Email
10. Application Method

============================================================
⚠️ JOB RISK DETECTOR
============================================================

Check whether the JD mentions:

1. Bond
2. Agreement
3. Service Agreement
4. Contract Period
5. Lock-in Period
6. Training Agreement
7. Penalty for Leaving
8. Security Deposit
9. Salary Deduction
10. Unpaid Training
11. Long Notice Period
12. Work From Office Requirement
13. Relocation Requirement
14. Travel Requirement
15. Shift Requirement
16. Night Shift
17. Weekend Work
18. Overtime
19. Sales / Target Requirement
20. Other important conditions

For every detected condition provide:

Condition:
Status:
Details:
Risk Level:
What Candidate Should Check:

Risk Level:

🟢 Low
🟡 Medium
🔴 High

============================================================
💰 SALARY ANALYSIS
============================================================

If salary is mentioned:

- Show exact salary / CTC.
- Explain fixed or variable if mentioned.
- Mention incentives if mentioned.

If salary is NOT mentioned:

⚠️ Salary Not Mentioned

============================================================
📍 LOCATION ANALYSIS
============================================================

Show:

- Work Location
- Work Mode
- Relocation Requirement
- Travel Requirement

Only if mentioned.

============================================================
📜 AGREEMENT / BOND ANALYSIS
============================================================

If bond/agreement exists, explain:

- Duration
- Conditions
- Penalty
- Exit conditions
- Amount
- Restrictions

If none is mentioned:

✅ No Bond/Agreement Mentioned

============================================================
📧 APPLICATION DETAILS
============================================================

If an email is mentioned:

Show exact email.

If application instructions are mentioned:

Explain exactly how to apply.

============================================================
🚨 FINAL JOB RISK SUMMARY
============================================================

Give:

Overall Risk Level:
Major Concern:
Things to Verify Before Applying:
Positive Points:
Final Recommendation:

IMPORTANT:

- ONLY use information explicitly present in the JD.
- Never invent salary.
- Never invent company details.
- Never invent a bond.
- Never claim a job is fraudulent.
- If something is not mentioned say:
  "Not mentioned in the job description."

"""

    response = llm.invoke(
        prompt
    )

    return response.content


# ============================================================
# 16. BASIC INTERVIEW QUESTIONS
# ============================================================

def generate_questions(
    job_role,
    job_description=""
):

    prompt = f"""

You are a friendly interviewer conducting a normal
entry-level / fresher interview.

JOB ROLE:

{job_role}

JOB DESCRIPTION:

{job_description}

The candidate is a fresher.

Generate simple and commonly asked questions.

DO NOT ask:

- Advanced questions
- Expert-level questions
- Research-level questions
- Complex system design
- Very difficult coding questions

Generate:

5 Basic Technical Questions
2 Simple Project Questions
2 Simple Behavioral Questions
1 Simple HR Question

Requirements:

- Easy to moderate difficulty
- Suitable for fresh graduates
- Commonly asked interview questions
- Short and clear
- Relevant to the job

Do not provide answers.

"""

    response = llm.invoke(
        prompt
    )

    return response.content


# ============================================================
# 17. PERSONALIZED QUESTIONS
# ============================================================

def generate_personalized_questions(
    resume_text,
    job_description
):

    prompt = f"""

You are a friendly interviewer conducting a normal
entry-level / fresher interview.

CANDIDATE RESUME:

{resume_text}

JOB DESCRIPTION:

{job_description}

Generate personalized interview questions.

The candidate is a fresher.

DO NOT ask:

- Advanced system design
- Research-level questions
- Very difficult algorithms
- Complex mathematical questions
- Expert-level architecture questions

Generate:

5 Basic Technical Questions
3 Simple Project Questions
2 Questions about Resume Skills
2 Questions related to Job Description
2 Simple Behavioral Questions
1 Simple HR Question

Requirements:

- Easy to moderate difficulty
- Suitable for a fresher
- Based on actual resume
- Relevant to job description
- Short and clear

Do not provide answers.

"""

    response = llm.invoke(
        prompt
    )

    return response.content


# ============================================================
# 18. ANSWER EVALUATION
# ============================================================

def evaluate_answer(
    question,
    user_answer
):

    ideal_prompt = f"""

Provide a simple and technically correct ideal answer
for this interview question.

QUESTION:

{question}

The answer should be suitable for a fresher.

"""

    ideal_answer = llm.invoke(
        ideal_prompt
    ).content

    user_embedding = embedding_model.encode(
        [user_answer]
    )

    ideal_embedding = embedding_model.encode(
        [ideal_answer]
    )

    similarity_score = cosine_similarity(
        user_embedding,
        ideal_embedding
    )[0][0]

    similarity_score = max(
        0,
        min(
            1,
            similarity_score
        )
    )

    final_score = round(
        similarity_score * 10,
        2
    )

    feedback_prompt = f"""

You are a friendly technical interviewer.

QUESTION:

{question}

USER ANSWER:

{user_answer}

IDEAL ANSWER:

{ideal_answer}

Provide:

1. Score out of 10
2. Strengths
3. Weaknesses
4. Missing Points
5. Answer Accuracy
6. Simple Improvement Suggestions
7. Better Sample Answer

Keep the feedback easy to understand.

"""

    feedback = llm.invoke(
        feedback_prompt
    ).content

    return (
        final_score,
        feedback
    )


# ============================================================
# 19. EXTRACT RESUME
# ============================================================

if uploaded_file:

    resume_text = extract_resume_text(
        uploaded_file
    )

else:

    resume_text = ""


# ============================================================
# 20. HERO SECTION
# ============================================================

st.markdown(
    '<div class="hero-box">'
    '<div class="hero-small">'
    'AI CAREER INTELLIGENCE PLATFORM'
    '</div>'
    '<div class="hero-title">'
    'Build a better career with <span>AI</span>'
    '</div>'
    '<div class="hero-description">'
    'Analyze your resume, compare it with job requirements, '
    'identify missing skills and prepare for interviews '
    'with personalized AI guidance.'
    '</div>'
    '<div class="feature-row">'
    '<div class="feature-badge">📄 Resume Analysis</div>'
    '<div class="feature-badge">🎯 Job Matching</div>'
    '<div class="feature-badge">🧩 Skill Gap Detection</div>'
    '<div class="feature-badge">📝 Resume Improvement</div>'
    '<div class="feature-badge">📂 Project Analyzer</div>'
    '<div class="feature-badge">⚠️ JD Risk Detector</div>'
    '<div class="feature-badge">🎤 Interview Prep</div>'
    '</div>'
    '</div>',
    unsafe_allow_html=True
)


# ============================================================
# 21. DASHBOARD
# ============================================================

if uploaded_file:

    st.markdown(
        '<div class="dashboard-heading">'
        '<div class="section-title">'
        'Welcome to your <span>Career Dashboard</span>'
        '</div>'
        '<div class="section-description">'
        'Everything you need to understand your job fit '
        'and prepare for your interview.'
        '</div>'
        '</div>',
        unsafe_allow_html=True
    )

    col1, col2, col3 = st.columns(
        3,
        gap="medium"
    )

    with col1:

        st.markdown(
            '<div class="dashboard-card">'
            '<div>'
            '<div class="dashboard-icon">📄</div>'
            '<div class="dashboard-title">'
            'Resume'
            '</div>'
            '</div>'
            f'<div class="dashboard-value">'
            f'{uploaded_file.name}'
            '</div>'
            '</div>',
            unsafe_allow_html=True
        )

    with col2:

        st.markdown(
            '<div class="dashboard-card">'
            '<div>'
            '<div class="dashboard-icon">💼</div>'
            '<div class="dashboard-title">'
            'Target Role'
            '</div>'
            '</div>'
            f'<div class="dashboard-value">'
            f'{job_role if job_role else "Not selected"}'
            '</div>'
            '</div>',
            unsafe_allow_html=True
        )

    with col3:

        st.markdown(
            '<div class="dashboard-card">'
            '<div>'
            '<div class="dashboard-icon">📋</div>'
            '<div class="dashboard-title">'
            'Job Description'
            '</div>'
            '</div>'
            f'<div class="dashboard-value">'
            f'{"Added ✓" if job_description else "Not added"}'
            '</div>'
            '</div>',
            unsafe_allow_html=True
        )

    st.markdown(
        '<div class="dashboard-bottom-space"></div>',
        unsafe_allow_html=True
    )

else:

    st.markdown(
        '<div class="welcome-card">'
        '<div class="welcome-title">'
        '👋 Welcome to CareerAI'
        '</div>'
        '<div class="welcome-text">'
        'Upload your resume from the sidebar to get started. '
        'CareerAI can analyze your resume, compare it with '
        'a job description, identify missing skills and help '
        'you prepare for interviews.'
        '</div>'
        '</div>',
        unsafe_allow_html=True
    )


# ============================================================
# 22. TABS
# ============================================================

tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs(
    [
        "📊 Resume Analysis",
        "🎯 Job Compatibility",
        "📝 Resume Improvement",
        "📂 Resume Project Analyzer",
        "⚠️ JD Risk Detector",
        "🎤 Interview Preparation",
        "🧠 Practice Interview"
    ]
)


# ============================================================
# TAB 1 - RESUME ANALYSIS
# ============================================================

with tab1:

    st.markdown(
        '<div class="section-title">'
        '📊 Resume <span>Analysis</span>'
        '</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="section-description">'
        'Get an AI-powered evaluation of your resume '
        'for your target role.'
        '</div>',
        unsafe_allow_html=True
    )

    if not uploaded_file:

        st.info(
            "📄 Upload your resume from the sidebar."
        )

    elif not job_role:

        st.info(
            "💼 Enter your target job role."
        )

    else:

        if st.button(
            "🔍 Analyze My Resume",
            key="resume_analysis_button"
        ):

            with st.spinner(
                "🤖 AI is analyzing your resume..."
            ):

                st.session_state.resume_analysis = (
                    evaluate_resume(
                        resume_text,
                        job_role
                    )
                )

        if st.session_state.resume_analysis:

            st.markdown(
                '<div class="card">'
                '<div class="card-title">'
                '🤖 AI Resume Evaluation'
                '</div>'
                '</div>',
                unsafe_allow_html=True
            )

            st.write(
                st.session_state.resume_analysis
            )


# ============================================================
# TAB 2 - JOB COMPATIBILITY
# ============================================================

with tab2:

    st.markdown(
        '<div class="section-title">'
        '🎯 Job <span>Compatibility</span>'
        '</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="section-description">'
        'Compare your resume with the specific job description '
        'and identify missing skills.'
        '</div>',
        unsafe_allow_html=True
    )

    if not uploaded_file:

        st.info(
            "📄 Upload your resume first."
        )

    elif not job_description:

        st.info(
            "📋 Paste the job description in the sidebar."
        )

    else:

        if st.button(
            "🚀 Analyze Job Compatibility",
            key="job_analysis_button"
        ):

            with st.spinner(
                "🔎 Comparing resume with job description..."
            ):

                st.session_state.job_analysis = (
                    compare_resume_with_job(
                        resume_text,
                        job_description
                    )
                )

        if st.session_state.job_analysis:

            (
                semantic_score,
                resume_skills,
                job_skills,
                missing_skills,
                comparison_result
            ) = st.session_state.job_analysis


            score_col1, score_col2, score_col3 = st.columns(
                3,
                gap="medium"
            )

            with score_col1:

                st.markdown(
                    '<div class="score-card">'
                    f'<div class="score-number">'
                    f'{semantic_score}%'
                    '</div>'
                    '<div class="score-label">'
                    'Resume–Job Similarity'
                    '</div>'
                    '</div>',
                    unsafe_allow_html=True
                )

            with score_col2:

                st.markdown(
                    '<div class="score-card">'
                    f'<div class="score-number">'
                    f'{len(resume_skills)}'
                    '</div>'
                    '<div class="score-label">'
                    'Resume Skills Found'
                    '</div>'
                    '</div>',
                    unsafe_allow_html=True
                )

            with score_col3:

                st.markdown(
                    '<div class="score-card">'
                    f'<div class="score-number">'
                    f'{len(missing_skills)}'
                    '</div>'
                    '<div class="score-label">'
                    'Missing Skills'
                    '</div>'
                    '</div>',
                    unsafe_allow_html=True
                )


            st.write("")


            col1, col2 = st.columns(
                2,
                gap="medium"
            )

            with col1:

                st.markdown(
                    "### 📄 Skills in Your Resume"
                )

                skill_html = ""

                for skill in resume_skills:

                    skill_html += (
                        f'<span class="skill">'
                        f'{skill}'
                        f'</span>'
                    )

                if skill_html:

                    st.markdown(
                        f'<div class="skill-container">'
                        f'{skill_html}'
                        f'</div>',
                        unsafe_allow_html=True
                    )

                else:

                    st.warning(
                        "No skills detected."
                    )


            with col2:

                st.markdown(
                    "### 💼 Skills Required by Job"
                )

                skill_html = ""

                for skill in job_skills:

                    skill_html += (
                        f'<span class="skill">'
                        f'{skill}'
                        f'</span>'
                    )

                if skill_html:

                    st.markdown(
                        f'<div class="skill-container">'
                        f'{skill_html}'
                        f'</div>',
                        unsafe_allow_html=True
                    )

                else:

                    st.warning(
                        "No skills detected."
                    )


            st.write("")


            st.markdown(
                "### ❌ Missing Skills for This Job"
            )

            if missing_skills:

                st.markdown(
                    '<div class="info-box">'
                    '⚠️ These skills are mentioned in the '
                    'job description but were not found '
                    'in your resume.'
                    '</div>',
                    unsafe_allow_html=True
                )

                missing_html = ""

                for skill in missing_skills:

                    missing_html += (
                        f'<span class="missing-skill">'
                        f'❌ {skill}'
                        f'</span>'
                    )

                st.markdown(
                    f'<div class="skill-container">'
                    f'{missing_html}'
                    f'</div>',
                    unsafe_allow_html=True
                )

            else:

                st.success(
                    "🎉 No major missing skills detected!"
                )


            st.write("")


            st.markdown(
                "### 🤖 AI Job Compatibility Analysis"
            )

            st.write(
                comparison_result
            )


# ============================================================
# TAB 3 - RESUME IMPROVEMENT
# ============================================================

with tab3:

    st.markdown(
        '<div class="section-title">'
        '📝 Resume <span>Improvement</span>'
        '</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="section-description">'
        'Get resume improvement suggestions based specifically '
        'on the job requirements, missing skills and missing '
        'capabilities for the selected job.'
        '</div>',
        unsafe_allow_html=True
    )

    if not uploaded_file:

        st.info(
            "📄 Upload your resume first."
        )

    elif not job_description:

        st.info(
            "📋 Paste the job description first."
        )

    elif not job_role:

        st.info(
            "💼 Enter your target job role first."
        )

    else:

        if st.button(
            "📝 Generate Resume Improvements",
            key="resume_improvement_button"
        ):

            with st.spinner(
                "🤖 AI is finding job-specific resume improvements..."
            ):

                # Run Job Compatibility if it has not already
                # been generated.

                if st.session_state.job_analysis:

                    (
                        semantic_score,
                        resume_skills,
                        job_skills,
                        missing_skills,
                        comparison_result
                    ) = st.session_state.job_analysis

                else:

                    (
                        semantic_score,
                        resume_skills,
                        job_skills,
                        missing_skills,
                        comparison_result
                    ) = compare_resume_with_job(
                        resume_text,
                        job_description
                    )

                    st.session_state.job_analysis = (
                        semantic_score,
                        resume_skills,
                        job_skills,
                        missing_skills,
                        comparison_result
                    )

                st.session_state.resume_improvement = (
                    generate_resume_improvement(
                        resume_text,
                        job_description,
                        job_role,
                        missing_skills
                    )
                )


        if st.session_state.resume_improvement:

            st.markdown(
                '<div class="generated-box">',
                unsafe_allow_html=True
            )

            st.markdown(
                st.session_state.resume_improvement
            )

            st.markdown(
                '</div>',
                unsafe_allow_html=True
            )


# ============================================================
# TAB 4 - RESUME PROJECT ANALYZER
# ============================================================

with tab4:

    st.markdown(
        '<div class="section-title">'
        '📂 Resume <span>Project Analyzer</span>'
        '</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="section-description">'
        'Analyze the projects in your resume and understand '
        'how strongly they match the target job.'
        '</div>',
        unsafe_allow_html=True
    )

    if not uploaded_file:

        st.info(
            "📄 Upload your resume first."
        )

    elif not job_role:

        st.info(
            "💼 Enter your target job role first."
        )

    elif not job_description:

        st.info(
            "📋 Paste the job description first."
        )

    else:

        if st.button(
            "📂 Analyze My Resume Projects",
            key="project_analysis_button"
        ):

            with st.spinner(
                "🤖 AI is analyzing your projects..."
            ):

                st.session_state.project_analysis = (
                    analyze_resume_projects(
                        resume_text,
                        job_role,
                        job_description
                    )
                )

        if st.session_state.project_analysis:

            st.markdown(
                '<div class="generated-box">',
                unsafe_allow_html=True
            )

            st.markdown(
                st.session_state.project_analysis
            )

            st.markdown(
                '</div>',
                unsafe_allow_html=True
            )


# ============================================================
# TAB 5 - JOB DESCRIPTION RISK DETECTOR
# ============================================================

with tab5:

    st.markdown(
        '<div class="section-title">'
        '⚠️ Job Description <span>Risk Detector</span>'
        '</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="section-description">'
        'Check important job conditions such as salary, '
        'location, work mode, bond, agreement, shifts and '
        'application details before applying.'
        '</div>',
        unsafe_allow_html=True
    )

    if not job_description:

        st.info(
            "📋 Paste the job description in the sidebar."
        )

    else:

        if st.button(
            "⚠️ Analyze Job Risks",
            key="job_risk_button"
        ):

            with st.spinner(
                "🔎 Checking job description for important conditions..."
            ):

                st.session_state.job_risk_analysis = (
                    analyze_job_description_risks(
                        job_description
                    )
                )

        if st.session_state.job_risk_analysis:

            st.markdown(
                '<div class="generated-box">',
                unsafe_allow_html=True
            )

            st.markdown(
                st.session_state.job_risk_analysis
            )

            st.markdown(
                '</div>',
                unsafe_allow_html=True
            )


# ============================================================
# TAB 6 - INTERVIEW PREPARATION
# ============================================================

with tab6:

    st.markdown(
        '<div class="section-title">'
        '🎤 Interview <span>Preparation</span>'
        '</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="section-description">'
        'Practice simple and commonly asked fresher-level '
        'interview questions.'
        '</div>',
        unsafe_allow_html=True
    )


    basic_col, personalized_col = st.columns(
        2,
        gap="medium"
    )


    with basic_col:

        st.markdown(
            '<div class="question-card">'
            '<div class="question-card-title">'
            '📝 Basic Interview Questions'
            '</div>'
            '<div class="question-card-description">'
            'Simple technical, project, behavioral and HR '
            'questions commonly asked in fresher interviews.'
            '</div>'
            '</div>',
            unsafe_allow_html=True
        )


        if uploaded_file and job_role:

            if st.button(
                "Generate Basic Questions",
                key="generate_basic_questions"
            ):

                with st.spinner(
                    "💬 Generating basic questions..."
                ):

                    st.session_state.basic_questions = (
                        generate_questions(
                            job_role,
                            job_description
                        )
                    )


            if st.session_state.basic_questions:

                st.markdown(
                    '<div class="generated-box">',
                    unsafe_allow_html=True
                )

                st.markdown(
                    st.session_state.basic_questions
                )

                st.markdown(
                    '</div>',
                    unsafe_allow_html=True
                )

        else:

            st.info(
                "Upload resume and enter your target role."
            )


    with personalized_col:

        st.markdown(
            '<div class="question-card">'
            '<div class="question-card-title">'
            '🎯 Personalized Questions'
            '</div>'
            '<div class="question-card-description">'
            'Questions generated specifically from your '
            'resume and job description.'
            '</div>'
            '</div>',
            unsafe_allow_html=True
        )


        if uploaded_file and job_description:

            if st.button(
                "Generate Personalized Questions",
                key="generate_personalized_questions"
            ):

                with st.spinner(
                    "🧠 Creating personalized questions..."
                ):

                    st.session_state.personalized_questions = (
                        generate_personalized_questions(
                            resume_text,
                            job_description
                        )
                    )


            if st.session_state.personalized_questions:

                st.markdown(
                    '<div class="generated-box">',
                    unsafe_allow_html=True
                )

                st.markdown(
                    st.session_state.personalized_questions
                )

                st.markdown(
                    '</div>',
                    unsafe_allow_html=True
                )

        else:

            st.info(
                "Upload resume and paste job description."
            )


# ============================================================
# TAB 7 - PRACTICE INTERVIEW
# ============================================================

with tab7:

    st.markdown(
        '<div class="section-title">'
        '🧠 Practice <span>Interview</span>'
        '</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="section-description">'
        'Enter an interview question and your answer '
        'to receive AI-powered feedback.'
        '</div>',
        unsafe_allow_html=True
    )


    question_input = st.text_input(
        "Interview Question",
        placeholder="Example: What is Machine Learning?"
    )


    answer_input = st.text_area(
        "Your Answer",
        height=180,
        placeholder="Type your interview answer here..."
    )


    if st.button(
        "🎯 Evaluate My Answer",
        key="evaluate_answer_button"
    ):

        if question_input and answer_input:

            with st.spinner(
                "🤖 AI is evaluating your answer..."
            ):

                score, feedback = evaluate_answer(
                    question_input,
                    answer_input
                )


            col1, col2 = st.columns(
                2,
                gap="medium"
            )


            with col1:

                st.markdown(
                    '<div class="score-card">'
                    f'<div class="score-number">'
                    f'{score}/10'
                    '</div>'
                    '<div class="score-label">'
                    'Answer Similarity Score'
                    '</div>'
                    '</div>',
                    unsafe_allow_html=True
                )


            with col2:

                st.markdown(
                    '<div class="score-card">'
                    '<div class="score-number">'
                    '🤖'
                    '</div>'
                    '<div class="score-label">'
                    'AI Evaluation'
                    '</div>'
                    '</div>',
                    unsafe_allow_html=True
                )


            st.write("")


            st.markdown(
                "### 📝 Detailed AI Feedback"
            )

            st.write(
                feedback
            )

        else:

            st.warning(
                "Please enter both the question "
                "and your answer."
            )


# ============================================================
# FOOTER
# ============================================================

st.markdown(
    '<div class="footer">'
    '🤖 CareerAI · AI-Powered Resume & Interview Assistant'
    '<br><br>'
    'Resume Analysis · Job Matching · Skill Gap Detection · '
    'Resume Improvement · Project Analyzer · '
    'JD Risk Detector · Interview Preparation'
    '</div>',
    unsafe_allow_html=True
)
